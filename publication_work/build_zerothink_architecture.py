from pathlib import Path
import re
from xml.sax.saxutils import escape

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, KeepTogether

ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "papers" / "zerothink-reproducible-architecture-1.0.md"
OUTDIR = ROOT / "publication_work" / "zerothink-architecture-1.0"
DOCX = OUTDIR / "ZeroThink-Reproducible-Architecture-1.0-Shafaet-Brady-Hussain.docx"
PDF = OUTDIR / "ZeroThink-Reproducible-Architecture-1.0-Shafaet-Brady-Hussain.pdf"
OUTDIR.mkdir(parents=True, exist_ok=True)

NAVY = RGBColor(11, 37, 69)
BLUE = RGBColor(46, 116, 181)
MUTED = RGBColor(89, 89, 89)

def set_font(run, name="Calibri", size=11, bold=False, italic=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color

def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tcMar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")

def shade(cell, fill):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shd)

def add_page_number(paragraph):
    run = paragraph.add_run()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    run._r.addnext(fld)

def inline_runs(p, text, base_size=11, base_color=None):
    parts = re.split(r"(`[^`]+`|\*\*[^*]+\*\*|\*[^*]+\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            r = p.add_run(part[1:-1]); set_font(r, "Consolas", base_size - .5, color=base_color)
        elif part.startswith("**") and part.endswith("**"):
            r = p.add_run(part[2:-2]); set_font(r, size=base_size, bold=True, color=base_color)
        elif part.startswith("*") and part.endswith("*"):
            r = p.add_run(part[1:-1]); set_font(r, size=base_size, italic=True, color=base_color)
        else:
            r = p.add_run(part); set_font(r, size=base_size, color=base_color)

def parse_markdown():
    lines = SRC.read_text(encoding="utf-8").splitlines()
    if lines and lines[0] == "---":
        end = lines.index("---", 1)
        lines = lines[end+1:]
    blocks, i = [], 0
    while i < len(lines):
        line = lines[i].rstrip()
        if not line:
            i += 1; continue
        if line.startswith("#"):
            m = re.match(r"^(#{1,3})\s+(.*)$", line)
            if m: blocks.append(("heading", len(m.group(1)), m.group(2))); i += 1; continue
        if line.startswith("|"):
            rows = []
            while i < len(lines) and lines[i].lstrip().startswith("|"):
                cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                if not all(re.fullmatch(r":?-{3,}:?", c or "") for c in cells): rows.append(cells)
                i += 1
            blocks.append(("table", rows)); continue
        if re.match(r"^\d+\.\s+", line) or line.startswith("- "):
            ordered = bool(re.match(r"^\d+\.\s+", line)); items=[]
            while i < len(lines):
                cur=lines[i].rstrip(); m=re.match(r"^\d+\.\s+(.*)$",cur) if ordered else re.match(r"^-\s+(.*)$",cur)
                if not m: break
                items.append(m.group(1)); i+=1
            blocks.append(("list", ordered, items)); continue
        para=[line]; i+=1
        while i < len(lines) and lines[i].strip() and not lines[i].startswith("#") and not lines[i].lstrip().startswith("|") and not re.match(r"^(\d+\.|-)\s+", lines[i]):
            para.append(lines[i].strip()); i+=1
        blocks.append(("paragraph", " ".join(para)))
    return blocks

def build_docx(blocks):
    doc=Document(); sec=doc.sections[0]
    sec.page_width=Inches(8.5); sec.page_height=Inches(11)
    sec.top_margin=sec.bottom_margin=sec.left_margin=sec.right_margin=Inches(1)
    sec.header_distance=sec.footer_distance=Inches(.492)
    normal=doc.styles["Normal"]
    normal.font.name="Calibri"; normal._element.rPr.rFonts.set(qn("w:ascii"),"Calibri"); normal.font.size=Pt(11)
    normal.paragraph_format.space_after=Pt(8); normal.paragraph_format.line_spacing=1.333
    for name,size,color,before,after in (("Heading 1",16,BLUE,18,10),("Heading 2",13,BLUE,12,6),("Heading 3",12,RGBColor(31,77,120),8,4)):
        s=doc.styles[name]; s.font.name="Calibri"; s._element.rPr.rFonts.set(qn("w:ascii"),"Calibri"); s.font.size=Pt(size); s.font.bold=True; s.font.color.rgb=color; s.paragraph_format.space_before=Pt(before); s.paragraph_format.space_after=Pt(after)
    header=sec.header.paragraphs[0]; header.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    set_font(header.add_run("ZEROTHINK 1.0 | REPRODUCIBLE ARCHITECTURE | PUBLIC PREPRINT"),size=8,color=MUTED)
    footer=sec.footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.CENTER
    set_font(footer.add_run("Shafaet Brady Hussain  |  "),size=8,color=MUTED); add_page_number(footer)
    first_h1=True
    for block in blocks:
        if block[0]=="heading":
            _,level,text=block
            if level==1 and first_h1:
                p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(72); p.paragraph_format.space_after=Pt(8)
                inline_runs(p,text,24,NAVY); p.runs[0].bold=True; first_h1=False
            elif level==2 and not any(x in text for x in ("Abstract","References","Appendix")) and len(doc.paragraphs)<5:
                p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(16); inline_runs(p,text,14,BLUE)
            else:
                doc.add_heading(text,level=min(level,3))
        elif block[0]=="paragraph":
            text=block[1]
            p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY; p.paragraph_format.space_after=Pt(8); p.paragraph_format.line_spacing=1.333
            inline_runs(p,text)
        elif block[0]=="list":
            _,ordered,items=block
            for item in items:
                p=doc.add_paragraph(style="List Number" if ordered else "List Bullet"); p.paragraph_format.left_indent=Inches(.375); p.paragraph_format.first_line_indent=Inches(-.194); p.paragraph_format.space_after=Pt(4); p.paragraph_format.line_spacing=1.208; inline_runs(p,item)
        elif block[0]=="table":
            rows=block[1]
            if not rows: continue
            cols=max(len(r) for r in rows); table=doc.add_table(rows=0,cols=cols); table.style="Table Grid"; table.alignment=WD_TABLE_ALIGNMENT.CENTER; table.autofit=False
            widths=[Inches(6.5/cols)]*cols
            for ri,row in enumerate(rows):
                cells=table.add_row().cells
                for ci in range(cols):
                    cells[ci].width=widths[ci]; cells[ci].vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.TOP; set_cell_margins(cells[ci])
                    if ri==0: shade(cells[ci],"F4F6F9")
                    p=cells[ci].paragraphs[0]; p.paragraph_format.space_after=Pt(0); p.paragraph_format.line_spacing=1.05
                    inline_runs(p,row[ci] if ci<len(row) else "",9)
                    if ri==0:
                        for run in p.runs: run.bold=True
            doc.add_paragraph().paragraph_format.space_after=Pt(2)
    doc.save(DOCX)

def clean_pdf_text(text):
    parts = re.split(r"(`[^`]+`|\*\*[^*]+\*\*|\*[^*]+\*)", text)
    out=[]
    for part in parts:
        if not part: continue
        if part.startswith("`") and part.endswith("`"):
            out.append('<font name="Courier">'+escape(part[1:-1])+'</font>')
        elif part.startswith("**") and part.endswith("**"):
            out.append('<b>'+escape(part[2:-2])+'</b>')
        elif part.startswith("*") and part.endswith("*"):
            out.append('<i>'+escape(part[1:-1])+'</i>')
        else: out.append(escape(part))
    return "".join(out)

def build_pdf(blocks):
    styles=getSampleStyleSheet()
    body=ParagraphStyle("BodyZ",parent=styles["BodyText"],fontName="Helvetica",fontSize=9.6,leading=13.2,alignment=TA_JUSTIFY,spaceAfter=7,textColor=colors.HexColor("#151A20"))
    h1=ParagraphStyle("H1Z",parent=styles["Heading1"],fontName="Helvetica-Bold",fontSize=15,leading=18,textColor=colors.HexColor("#2E74B5"),spaceBefore=15,spaceAfter=8)
    h2=ParagraphStyle("H2Z",parent=styles["Heading2"],fontName="Helvetica-Bold",fontSize=12,leading=15,textColor=colors.HexColor("#2E74B5"),spaceBefore=10,spaceAfter=5)
    h3=ParagraphStyle("H3Z",parent=styles["Heading3"],fontName="Helvetica-Bold",fontSize=10.5,leading=13,textColor=colors.HexColor("#1F4D78"),spaceBefore=8,spaceAfter=4)
    title=ParagraphStyle("TitleZ",parent=styles["Title"],fontName="Helvetica-Bold",fontSize=23,leading=27,textColor=colors.HexColor("#0B2545"),alignment=TA_CENTER,spaceAfter=12)
    subtitle=ParagraphStyle("SubZ",parent=body,fontName="Helvetica",fontSize=13,leading=17,textColor=colors.HexColor("#2E74B5"),alignment=TA_CENTER,spaceAfter=20)
    bullet=ParagraphStyle("BulletZ",parent=body,leftIndent=18,firstLineIndent=-10,spaceAfter=4)
    story=[]; first=True
    for block in blocks:
        if block[0]=="heading":
            _,level,text=block
            if level==1 and first:
                story += [Spacer(1,1.2*inch),Paragraph(clean_pdf_text(text),title)]; first=False
            elif level==2 and len(story)<5:
                story.append(Paragraph(clean_pdf_text(text),subtitle))
            else: story.append(Paragraph(clean_pdf_text(text),{1:h1,2:h2,3:h3}[min(level,3)]))
        elif block[0]=="paragraph": story.append(Paragraph(clean_pdf_text(block[1]),body))
        elif block[0]=="list":
            _,ordered,items=block
            for i,item in enumerate(items,1): story.append(Paragraph((f"{i}. " if ordered else "&#8226; ")+clean_pdf_text(item),bullet))
        elif block[0]=="table":
            rows=block[1]
            data=[[Paragraph(clean_pdf_text(c),ParagraphStyle("Cell",parent=body,fontSize=7.6,leading=9.4,spaceAfter=0,alignment=TA_LEFT)) for c in row] for row in rows]
            if data:
                t=Table(data,colWidths=[6.5*inch/len(data[0])]*len(data[0]),repeatRows=1,hAlign="LEFT")
                t.setStyle(TableStyle([("GRID",(0,0),(-1,-1),.4,colors.HexColor("#B9C3CF")),("BACKGROUND",(0,0),(-1,0),colors.HexColor("#F4F6F9")),("FONTNAME",(0,0),(-1,0),"Helvetica-Bold"),("VALIGN",(0,0),(-1,-1),"TOP"),("LEFTPADDING",(0,0),(-1,-1),5),("RIGHTPADDING",(0,0),(-1,-1),5),("TOPPADDING",(0,0),(-1,-1),4),("BOTTOMPADDING",(0,0),(-1,-1),4)])); story += [t,Spacer(1,6)]
    def footer(canvas,doc):
        canvas.saveState(); canvas.setFont("Helvetica",7.5); canvas.setFillColor(colors.HexColor("#595959")); canvas.drawString(inch,.55*inch,"ZEROTHINK 1.0 | REPRODUCIBLE ARCHITECTURE | PUBLIC PREPRINT"); canvas.drawRightString(7.5*inch,.55*inch,f"Shafaet Brady Hussain | {doc.page}"); canvas.restoreState()
    pdf=SimpleDocTemplate(str(PDF),pagesize=letter,rightMargin=inch,leftMargin=inch,topMargin=.85*inch,bottomMargin=.8*inch,title="ZeroThink 1.0: A Reproducible Architecture for Evidence-Gated AI Agent Services",author="Shafaet Brady Hussain",subject="Evidence-gated AI agent architecture, boundary testing, authenticated routing, provenance, and reproducibility")
    pdf.build(story,onFirstPage=footer,onLaterPages=footer)

blocks=parse_markdown(); build_docx(blocks); build_pdf(blocks)
print(DOCX); print(PDF)
